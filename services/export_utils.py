import os
import tempfile
from datetime import datetime
from typing import Tuple

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


def _timestamped(prefix: str, ext: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    tmpdir = tempfile.mkdtemp(prefix="legal_assistant_")
    return os.path.join(tmpdir, f"{prefix}_{ts}.{ext}")


def export_docx(text: str, title: str = "AI Legal Document") -> str:
    """Create a DOCX file from text and return the file path."""
    path = _timestamped("document", "docx")
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for para in text.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(path)
    return path


essentials = [
    ("Important Notice:", "This document is AI-generated and must be reviewed by a qualified attorney."),
]


def export_pdf(text: str, title: str = "AI Legal Document") -> str:
    """Create a simple PDF file from text and return the file path."""
    path = _timestamped("document", "pdf")
    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER

    y = height - 72
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, y, title)
    y -= 24

    c.setFont("Helvetica", 11)
    for line in text.splitlines():
        if y < 72:
            c.showPage()
            y = height - 72
            c.setFont("Helvetica", 11)
        c.drawString(72, y, line[:1000])  # safety trim
        y -= 14

    # Footer notice
    if y < 100:
        c.showPage()
        y = height - 72
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, y, essentials[0][0])
    y -= 14
    c.setFont("Helvetica", 11)
    c.drawString(72, y, essentials[0][1])

    c.save()
    return path
